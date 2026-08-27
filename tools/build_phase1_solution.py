from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement


ROOT = Path(r"C:\Users\admin\Documents\ChatGPT\Approval System")
OUT = ROOT / "审批流微服务_第一阶段技术方案_V0.1.docx"

FONT = "Microsoft YaHei"
FONT_EAST_ASIA = "微软雅黑"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
GREEN = "EAF6EE"
AMBER = "FFF6DD"
RED = "FCEBEC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def setup_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, num_fmt, text, left=720, hanging=360, bullet_font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), bullet_font or FONT)
        fonts.set(qn("w:hAnsi"), bullet_font or FONT)
        fonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    add_abstract(100, "bullet", "•", bullet_font="Arial")
    add_num(100, 100)
    add_abstract(101, "decimal", "%1.")
    add_num(101, 101)


def add_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Body Text")
    add_numbering(p, 100)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="Body Text")
    add_numbering(p, 101)
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Body Text")
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill, color)
    set_run_font(p.add_run(title + "  "), size=10.5, color=color, bold=True)
    set_run_font(p.add_run(body), size=10.5, color=INK)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(h)), size=font_size, color=DARK_BLUE, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, "F7F8FA", "D0D5DD")
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(INK)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("第 "), size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    r_pr.append(r_fonts)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    set_run_font(paragraph.add_run(" 页"), size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    body = doc.styles["Body Text"]
    set_style_font(body, 11, INK)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    set_style_font(title, 25, INK, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 13, MUTED, False)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Small Note" not in [s.name for s in doc.styles]:
        note = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(note, 9, MUTED)
        note.paragraph_format.space_after = Pt(4)
        note.paragraph_format.line_spacing = 1.05


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("通用审批平台｜第一阶段技术方案"), size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_field(p)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("技术方案 / PHASE 1"), size=10, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    set_run_font(p.add_run("通用审批流微服务"), size=25, color=INK, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    set_run_font(p.add_run("第一阶段：轻量级可用版本（MVP）技术方案"), size=13, color=MUTED)

    meta = [
        ("文档版本", "V0.1"),
        ("文档状态", "方案草案 / 可用于立项评审"),
        ("编制日期", "2026-08-15"),
        ("目标读者", "产品负责人、技术负责人、架构师、研发与测试团队"),
        ("建议周期", "10-12 周，按团队规模可调整"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}："), size=10.5, color=INK, bold=True)
        set_run_font(p.add_run(value), size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_callout(
        doc,
        "核心建议",
        "第一阶段采用“一个审批核心服务 + 一个 API 网关 + 一个关系型数据库”的轻量部署。DDD 用于守住代码边界，Dubbo Triple 用于服务契约，Flowable 用于流程状态机；不在首期引入 Kubernetes、OpenSearch、Seata、复杂脚本平台或多个细粒度微服务。",
        fill=LIGHT_BLUE,
    )


def add_contents(doc):
    add_heading(doc, "文档目录", 1)
    sections = [
        "1. 执行摘要与架构决策", "2. 背景、目标与成功标准", "3. 第一阶段范围", "4. 总体架构",
        "5. DDD 领域设计", "6. 核心业务流程", "7. 功能需求", "8. API 与集成设计",
        "9. 规则与自动化", "10. 数据模型与一致性", "11. 安全与多租户", "12. 可观测性与运维",
        "13. 技术选型与工程规范", "14. 测试策略", "15. 实施计划", "16. 验收标准",
        "17. 风险与应对", "18. 后续演进", "附录 A-C：接口示例、状态模型、参考资料",
    ]
    add_table(doc, ["章节", "评审重点"], [
        ("1-4", "为什么这样做、首期是否足够轻、部署边界是否合理"),
        ("5-10", "领域模型、流程行为、接口契约和数据一致性是否完整"),
        ("11-14", "安全、运维、技术栈和质量保障是否达到可上线要求"),
        ("15-18", "计划、验收、风险和演进路径是否可执行"),
    ], [1600, 7760])
    p = doc.add_paragraph(style="Small Note")
    set_run_font(p.add_run("完整章节：" + "；".join(sections)), size=9, color=MUTED)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    setup_numbering(doc)
    add_title_block(doc)
    add_contents(doc)

    add_heading(doc, "1. 执行摘要与架构决策", 1)
    add_body(doc, "本项目建设一个面向多业务系统的通用审批服务。外部系统可以通过 API 发起、查询和管理审批；业务管理员可以配置流程模板、表单和规则；审批人员通过统一工作台处理待办。第一阶段优先验证通用性、稳定性和接入体验，而不是一次性建设完整 BPM 平台。")
    add_heading(doc, "1.1 已确定的关键决策", 2)
    add_table(doc, ["决策项", "第一阶段选择", "理由"], [
        ("部署形态", "approval-gateway + approval-center", "保持轻量；核心服务内部按 DDD 模块化，达到拆分条件后再独立部署"),
        ("内部通信", "Dubbo 3.x Triple", "满足 Java RPC、HTTP/2 和未来多语言扩展；服务契约与领域实现分离"),
        ("外部接口", "REST/JSON + OpenAPI + Webhook", "降低业务系统接入门槛，不要求外部系统使用 Dubbo"),
        ("流程引擎", "Flowable，封装为基础设施适配器", "复用 BPMN/任务/定时器能力，避免首期自研流程状态机"),
        ("规则", "CEL + 审批矩阵；复杂脚本延期", "覆盖大多数路由和自动审批，同时控制安全风险"),
        ("数据", "PostgreSQL 或 MySQL 单库起步", "审批是长事务，采用本地事务、审计日志和 Outbox，而不是全局 2PC"),
        ("AI", "预留接口；首期可选摘要试点", "AI 不进入最终确定性决策链，不阻塞 MVP 上线"),
    ], [1700, 2900, 4760], font_size=8.8)
    add_callout(doc, "架构边界", "DDD 限界上下文不等于微服务。第一阶段先在一个核心服务中实现清晰模块边界，只有出现独立扩容、故障隔离、团队所有权或发布节奏需求时才拆分。", fill=AMBER, color="7A5A00")

    add_heading(doc, "2. 背景、目标与成功标准", 1)
    add_heading(doc, "2.1 建设背景", 2)
    add_body(doc, "企业中的采购、费用、合同、权限、发布、数据申请等场景都需要审批。若每个业务系统自行实现，会重复建设审批状态、组织解析、通知、审计和超时处理，并形成互不兼容的数据与接口。本项目通过独立审批平台沉淀这些共性能力。")
    add_heading(doc, "2.2 第一阶段业务目标", 2)
    for item in [
        "外部系统能够在 1-2 个工作日内完成 API 接入和联调。",
        "管理员能够创建并发布通用审批模板、动态表单和条件规则。",
        "审批人能够处理待办、已办、同意、拒绝、转交和撤回等核心动作。",
        "系统能够可靠记录每一次状态变化，并将结果回调给外部系统。",
        "架构支持后续拆分微服务、接入 MQ、增强 AI 和企业身份体系，不需要推翻重做。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.3 建议量化指标", 2)
    add_table(doc, ["指标", "第一阶段目标", "测量方式"], [
        ("API 可用性", "月度不低于 99.9%（单区域）", "网关与核心接口 SLI"),
        ("启动审批延迟", "P95 < 500 ms，不含外部组织接口", "链路指标"),
        ("审批动作延迟", "P95 < 300 ms", "任务完成接口指标"),
        ("幂等正确性", "重复请求不产生重复实例/重复审批", "自动化与故障注入测试"),
        ("Webhook 交付", "最终成功率 ≥ 99.9%，支持重试和重放", "投递审计表"),
        ("审计完整性", "所有管理与审批动作 100% 留痕", "审计核对测试"),
    ], [1900, 2900, 4560])

    add_heading(doc, "3. 第一阶段范围", 1)
    add_heading(doc, "3.1 范围内", 2)
    for item in [
        "多租户基础隔离、应用凭证、用户/角色/部门适配。",
        "流程模板草稿、校验、发布和不可变版本。",
        "JSON Schema 动态表单、字段校验和节点级只读/可见控制。",
        "串行、并行、条件分支、或签/会签、抄送、超时提醒。",
        "发起、查询、同意、拒绝、撤回、转交、终止和审批意见。",
        "待办、已办、我发起的审批和管理端实例查询。",
        "CEL 条件、审批矩阵和低风险确定性自动审批。",
        "REST API、Dubbo API、Webhook、签名、幂等、重试和审计。",
        "基础监控、结构化日志、Trace、健康检查和备份恢复方案。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.2 明确不在第一阶段", 2)
    add_table(doc, ["延期能力", "原因 / 后续入口"], [
        ("流程实例跨版本迁移", "实现复杂且风险高；首期保证旧实例继续使用旧版本"),
        ("任意 Java/Groovy/Python 脚本", "存在远程代码执行风险；后续建设独立 Sandbox Runner"),
        ("流程挖掘与高级 BI", "首期保留事件和指标，第二阶段接入 OpenSearch/分析平台"),
        ("跨区域双活与单元化", "先完成单区域高可用，达到规模阈值后演进"),
        ("完整低代码页面搭建器", "首期支持基础表单配置，不建设通用低代码平台"),
        ("AI 自动做高风险最终决定", "首期 AI 仅做辅助，自动审批必须由确定性规则控制"),
    ], [3000, 6360])

    add_heading(doc, "4. 总体架构", 1)
    add_heading(doc, "4.1 逻辑架构", 2)
    add_table(doc, ["层级", "组件", "职责"], [
        ("接入层", "approval-gateway", "REST API、OAuth2/API Key、租户识别、限流、OpenAPI、管理页面 BFF"),
        ("应用层", "approval-center.application", "用例编排、事务边界、命令/查询、权限校验、DTO 转换"),
        ("领域层", "definition/runtime/task/rule/audit", "聚合、领域服务、领域事件和业务不变量"),
        ("基础设施层", "Flowable Adapter / Persistence / Webhook", "流程引擎、数据库、Outbox、组织适配、外部回调"),
        ("异步执行", "worker profile（可同进程或独立）", "Webhook、通知、超时扫描、Outbox 发布和失败重试"),
    ], [1500, 3000, 4860])

    add_heading(doc, "4.2 第一阶段部署拓扑", 2)
    add_callout(doc, "最小生产拓扑", "2 个无状态应用实例组（gateway、center）+ 1 个主备数据库。Redis、对象存储和 MQ 按场景启用；开发环境可用 Docker Compose 单机运行。", fill=GREEN, color="25633C")
    add_code_block(doc, """External Systems / Admin UI
              |
      REST / OpenAPI / Webhook
              v
      +-------------------+
      | approval-gateway  |
      +-------------------+
              | Dubbo Triple
              v
      +-------------------+
      | approval-center   |
      | DDD modules       |
      | Flowable adapter  |
      | outbox + worker   |
      +-------------------+
          |        |       \\
      PostgreSQL  Redis*   Object Storage*

* optional in the smallest deployment""")
    add_heading(doc, "4.3 调用与事件原则", 2)
    for item in [
        "需要即时结果的命令使用同步调用；跨系统通知、审计派生和报表更新使用异步事件。",
        "approval-gateway 不包含审批业务规则，只处理协议、安全和请求上下文。",
        "流程引擎不是领域模型；Flowable ID 仅保存在适配映射中，不作为公共 API 主标识。",
        "所有外部副作用在本地事务提交后执行，失败可以重试，不回滚已完成的人工审批事实。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. DDD 领域设计", 1)
    add_heading(doc, "5.1 限界上下文", 2)
    add_table(doc, ["上下文", "职责", "核心聚合 / 实体"], [
        ("Definition", "模板、流程定义、表单、版本、发布", "ApprovalTemplate、ProcessDefinition、FormSchema"),
        ("Runtime", "创建与推进审批实例、维护业务状态", "ApprovalInstance、ExecutionSnapshot"),
        ("Task", "待办分配、认领、转交、完成、会签统计", "ApprovalTask、Delegation"),
        ("Decision", "条件路由、审批矩阵、自动审批", "RuleSet、DecisionTable、DecisionResult"),
        ("Identity Adapter", "解析用户、部门、角色、直属主管", "PrincipalRef、GroupRef、ResolverPolicy"),
        ("Integration", "应用、Webhook、业务回调", "ClientApplication、WebhookSubscription、Delivery"),
        ("Audit", "不可变动作记录与证据", "AuditRecord、EvidenceRef"),
    ], [1750, 3600, 4010], font_size=8.8)
    add_heading(doc, "5.2 聚合边界与不变量", 2)
    for item in [
        "ApprovalInstance 是审批运行事实的聚合根；同一 externalBusinessKey 在租户和应用范围内唯一。",
        "ApprovalTask 独立聚合，使用 optimisticVersion 防止多人并发完成同一任务。",
        "已发布 ProcessDefinition 不允许原地修改；任何变更产生新版本。",
        "实例启动时固定 definitionVersion、formVersion、ruleVersion 和关键变量快照。",
        "审批动作必须同时写入业务状态、动作记录与 Outbox，三者处于同一本地事务。",
        "任何自动审批必须记录命中的规则、输入摘要、规则版本和结果。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "5.3 模块目录建议", 2)
    add_code_block(doc, """approval-center
├─ interfaces       # Dubbo、管理 REST、消息消费者
├─ application      # command、query、use case、DTO
├─ domain
│  ├─ definition
│  ├─ runtime
│  ├─ task
│  ├─ decision
│  ├─ integration
│  └─ audit
└─ infrastructure   # persistence、flowable、outbox、identity、webhook""")

    add_heading(doc, "6. 核心业务流程", 1)
    add_heading(doc, "6.1 模板发布", 2)
    for item in [
        "管理员创建模板草稿并配置表单、流程节点、审批人解析策略和规则。",
        "系统执行静态校验：入口/出口、不可达节点、循环、审批人为空、表达式类型和变量依赖。",
        "管理员在测试模式下使用样例数据模拟路径，并查看命中规则和预期审批人。",
        "发布后生成不可变版本；新实例使用新版本，存量实例保持原版本。",
    ]:
        add_number(doc, item)
    add_heading(doc, "6.2 发起审批", 2)
    for item in [
        "网关鉴权、识别 tenantId/applicationId，并校验 Idempotency-Key。",
        "核心服务读取已发布定义，校验表单和业务变量。",
        "创建 ApprovalInstance 和初始审计记录，启动 Flowable 实例。",
        "解析首批审批人并创建任务；若规则自动通过，则记录决策并继续推进。",
        "同一事务写入 Outbox；提交后异步发送通知和审批已创建 Webhook。",
    ]:
        add_number(doc, item)
    add_heading(doc, "6.3 处理审批任务", 2)
    for item in [
        "校验当前用户是否为处理人、候选人或有效委托人。",
        "校验任务版本和状态，防止重复点击、并发提交和越权操作。",
        "执行同意/拒绝/转交等领域命令，保存意见和附件引用。",
        "更新会签统计并决定是否推进、等待或结束流程。",
        "提交事务后异步通知下一审批人和外部业务系统。",
    ]:
        add_number(doc, item)
    add_heading(doc, "6.4 Webhook 交付", 2)
    add_body(doc, "Webhook 事件由 Outbox 驱动。每次投递生成 deliveryId，使用 HMAC 签名，携带事件 ID、事件时间、租户、业务键和审批快照。失败按指数退避重试，超过阈值进入失败队列，管理员可以查询并重放。接收方必须按 eventId 幂等。")

    add_heading(doc, "7. 功能需求", 1)
    add_heading(doc, "7.1 流程模板与表单", 2)
    add_table(doc, ["能力", "首期行为", "优先级"], [
        ("模板生命周期", "草稿、校验、发布、停用、复制、版本查看", "P0"),
        ("节点", "开始、人工审批、自动规则、并行/排他网关、抄送、结束", "P0"),
        ("审批人", "指定用户、角色、部门负责人、直属主管、表单字段", "P0"),
        ("会签", "全部通过、任一通过、比例通过；失败策略可配置", "P0"),
        ("动态表单", "JSON Schema + UI Schema、校验、附件、字段权限", "P0"),
        ("模拟", "样例变量下的路径、规则和审批人预览", "P1"),
    ], [2200, 5860, 1300])
    add_heading(doc, "7.2 审批实例与任务", 2)
    add_table(doc, ["能力", "首期行为", "优先级"], [
        ("实例", "创建、详情、时间线、撤回、终止、查询", "P0"),
        ("任务", "待办/已办、同意、拒绝、转交、审批意见", "P0"),
        ("超时", "提醒、升级到上级或指定角色；支持工作日历简版", "P1"),
        ("抄送", "创建抄送记录和通知，不参与流程推进", "P1"),
        ("委托", "按时间范围将任务委托给代理人", "P1"),
        ("批量审批", "仅对模板明确允许、低风险且相同动作的任务开放", "P1"),
    ], [2200, 5860, 1300])
    add_heading(doc, "7.3 管理与审计", 2)
    for item in [
        "租户、应用凭证、模板权限和操作员角色管理。",
        "按实例 ID、业务键、发起人、状态、模板、时间范围查询。",
        "流程时间线展示节点、处理人、动作、意见、时间、规则命中和回调结果。",
        "敏感字段按角色脱敏；审计导出需要单独权限。",
        "管理员修复动作需要双重确认、原因必填并产生高等级审计记录。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. API 与集成设计", 1)
    add_heading(doc, "8.1 API 设计原则", 2)
    for item in [
        "REST 资源命名稳定，公共接口采用 /v1 前缀；不暴露 Flowable 内部对象。",
        "写接口必须接受 Idempotency-Key；响应包含 requestId、instanceId 和当前版本。",
        "错误使用稳定业务错误码，不把 Java 异常、SQL 或内部堆栈返回客户端。",
        "列表使用 cursor 分页；查询接口允许 eventual consistency 的字段必须明确说明。",
        "Dubbo DTO 位于独立 contract 模块，禁止引用领域实体和持久化对象。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.2 首期公共接口", 2)
    add_table(doc, ["Method", "Path", "用途"], [
        ("POST", "/v1/approval-instances", "发起审批"),
        ("GET", "/v1/approval-instances/{id}", "查询审批详情与当前状态"),
        ("GET", "/v1/approval-instances", "按业务键、模板、状态等查询"),
        ("POST", "/v1/approval-instances/{id}/withdraw", "发起人撤回"),
        ("POST", "/v1/approval-instances/{id}/terminate", "管理员终止"),
        ("GET", "/v1/tasks", "待办、已办查询"),
        ("POST", "/v1/tasks/{id}/approve", "同意任务"),
        ("POST", "/v1/tasks/{id}/reject", "拒绝任务"),
        ("POST", "/v1/tasks/{id}/transfer", "转交任务"),
        ("POST", "/v1/webhook-subscriptions", "创建事件订阅"),
    ], [1200, 4300, 3860], font_size=9)
    add_heading(doc, "8.3 事件类型", 2)
    add_body(doc, "首期至少发布：approval.instance.created、approval.task.created、approval.task.completed、approval.instance.approved、approval.instance.rejected、approval.instance.withdrawn、approval.instance.terminated、webhook.delivery.failed。事件 Schema 必须版本化。")

    add_heading(doc, "9. 规则与自动化", 1)
    add_heading(doc, "9.1 规则分层", 2)
    add_table(doc, ["层级", "实现", "适用场景"], [
        ("条件表达式", "CEL", "金额、部门、风险等级、表单字段判断和简单数据转换"),
        ("审批矩阵", "结构化决策表", "金额区间 × 部门 × 业务类型映射审批链"),
        ("自动审批", "确定性规则 + 白名单", "低金额、低风险、可逆场景"),
        ("复杂脚本", "第一阶段不开放", "后续独立沙箱，限制网络、CPU、内存和执行时间"),
    ], [1800, 2800, 4760])
    add_heading(doc, "9.2 自动审批安全门", 2)
    for item in [
        "模板必须显式开启自动审批，并配置适用金额、风险级别和业务类型。",
        "规则执行只读取白名单变量；禁止访问数据库、网络、文件和系统环境。",
        "决策输出必须是结构化结果，不允许表达式直接执行副作用。",
        "记录 ruleId、ruleVersion、inputHash、result、reason 和 executionTime。",
        "规则异常时默认转人工，不得默认通过。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.3 AI 在第一阶段的边界", 2)
    add_callout(doc, "可选试点", "提供 AiAssistPort，仅试点“审批材料摘要”或“缺失信息提示”。AI 输出必须经过 JSON Schema 校验，并展示为建议；AI 不直接改变审批状态。", fill=AMBER, color="7A5A00")
    add_body(doc, "为后续演进预留 ai_execution 表或审计扩展字段，记录模型、Prompt 版本、输入来源、输出、置信度和人工是否采纳。涉及财务付款、权限授予和合同签署等高风险场景，默认要求人工最终决定。")

    add_heading(doc, "10. 数据模型与一致性", 1)
    add_heading(doc, "10.1 核心数据表", 2)
    add_table(doc, ["表 / 集合", "关键字段", "说明"], [
        ("approval_template", "tenant_id, template_key, status", "模板逻辑身份"),
        ("process_definition", "template_id, version, bpmn_xml, checksum", "不可变发布版本"),
        ("form_schema", "template_id, version, json_schema, ui_schema", "不可变表单版本"),
        ("approval_instance", "id, business_key, status, definition_version, lock_version", "审批当前状态"),
        ("approval_task", "id, instance_id, assignee, status, lock_version", "任务当前状态"),
        ("approval_action", "instance_id, task_id, actor, action, comment", "不可变动作时间线"),
        ("decision_record", "rule_id, rule_version, input_hash, result", "规则决策证据"),
        ("outbox_event", "event_id, aggregate_id, type, payload, publish_status", "可靠事件发布"),
        ("webhook_delivery", "event_id, endpoint, attempt, status, next_retry_at", "回调投递记录"),
        ("idempotency_record", "tenant_id, application_id, key, response_ref", "写请求防重"),
    ], [2100, 4060, 3200], font_size=8.5)
    add_heading(doc, "10.2 一致性策略", 2)
    for item in [
        "审批是持续数小时或数月的长流程，不使用数据库全局事务锁住全过程。",
        "单次业务动作使用本地 ACID：实例/任务状态、动作记录和 Outbox 原子提交。",
        "Outbox 发布器至少一次投递；消费者和 Webhook 接收方按 eventId 幂等。",
        "任务完成采用乐观锁和状态前置条件，例如仅 PENDING/CLAIMED 可完成。",
        "外部业务副作用失败后重试或补偿，不删除已形成的审批审计事实。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.3 数据保留与归档", 2)
    add_body(doc, "审批当前数据、审计记录和附件保留期应由租户策略控制。第一阶段至少实现逻辑删除限制、附件生命周期关联、数据库备份以及审计记录禁止普通管理员修改。大规模历史数据分区与归档在第二阶段实施。")

    add_heading(doc, "11. 安全与多租户", 1)
    add_heading(doc, "11.1 身份认证与授权", 2)
    for item in [
        "外部系统使用 OAuth2 Client Credentials，早期内部环境可兼容 API Key，但必须可轮换。",
        "管理端使用 OIDC/企业 SSO；审批权限同时检查用户身份、任务候选关系和租户边界。",
        "平台角色建议包括 TenantAdmin、TemplateDesigner、Auditor、Operator、Approver。",
        "模板和实例访问实施 RBAC + 资源归属校验；敏感场景增加 ABAC 条件。",
        "禁止仅依靠前端隐藏按钮实现授权，所有写操作必须在应用层再次校验。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "11.2 多租户隔离", 2)
    add_body(doc, "第一阶段采用共享数据库、共享 Schema、所有业务表强制 tenant_id 的模式。Repository 查询必须由 TenantContext 自动附加租户条件；唯一索引包含 tenant_id；缓存键、对象存储路径、日志和指标标签都必须带租户边界。高等级租户未来可迁移为独立 Schema 或独立数据库。")
    add_heading(doc, "11.3 安全控制清单", 2)
    add_table(doc, ["领域", "首期控制"], [
        ("传输", "TLS；内部 Dubbo Triple 启用身份认证或在受控网络内配合 mTLS"),
        ("存储", "数据库与对象存储加密；敏感字段应用级加密/脱敏"),
        ("Webhook", "HMAC 签名、时间戳、防重放、HTTPS、目标地址 SSRF 校验"),
        ("附件", "类型/大小限制、病毒扫描接口、临时授权下载"),
        ("审计", "登录、模板发布、审批、管理员修复、凭证变更全部留痕"),
        ("凭证", "通过 Secret Manager/KMS 管理，禁止写入代码和日志"),
    ], [2000, 7360])

    add_heading(doc, "12. 可观测性与运维", 1)
    add_heading(doc, "12.1 可观测性", 2)
    for item in [
        "使用 OpenTelemetry 统一 trace、metrics 和 logs，上下文携带 traceId、tenantId、applicationId、instanceId。",
        "核心指标：实例创建/完成数、任务处理延迟、待办积压、超时数、规则失败、Outbox 积压、Webhook 成功率。",
        "日志采用结构化 JSON，审批意见、表单敏感字段和 Token 不进入普通日志。",
        "健康检查区分 liveness、readiness，并检查数据库和流程引擎必要依赖。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.2 运维操作", 2)
    add_table(doc, ["操作", "要求"], [
        ("重放 Webhook", "按 deliveryId 重放；保留原事件并生成新尝试记录"),
        ("重新发布 Outbox", "仅对未发布/失败事件开放，不改变聚合状态"),
        ("终止卡死实例", "管理员权限、原因必填、审计记录、可选双人复核"),
        ("调整处理人", "走转交领域命令，不直接更新数据库"),
        ("数据修复", "禁止常态化 SQL 修改；使用受审计修复命令或一次性迁移"),
    ], [2700, 6660])
    add_heading(doc, "12.3 备份与恢复", 2)
    add_body(doc, "建议数据库每日全量备份并持续增量/WAL 归档，明确 RPO ≤ 15 分钟、RTO ≤ 2 小时作为首期目标。至少每季度执行一次恢复演练；对象存储附件和密钥备份纳入同一演练。")

    add_heading(doc, "13. 技术选型与工程规范", 1)
    add_heading(doc, "13.1 推荐技术栈", 2)
    add_table(doc, ["层面", "选择", "第一阶段备注"], [
        ("语言", "JDK 21", "统一 LTS 运行时"),
        ("应用", "Spring Boot 3.x", "Web、安全、配置与测试生态"),
        ("RPC", "Dubbo 3.x Triple", "内部契约；外部仍以 REST 为主"),
        ("流程", "Flowable", "嵌入 approval-center，通过 WorkflowEnginePort 隔离"),
        ("数据库", "PostgreSQL / MySQL", "二选一；生产使用主备/托管高可用"),
        ("访问层", "MyBatis 或 jOOQ", "显式 SQL，便于租户条件与性能治理"),
        ("规则", "CEL + 审批矩阵", "首期不开放任意脚本"),
        ("缓存", "Redis（可选）", "只做缓存、短期令牌/限流，不做审批事实源"),
        ("消息", "数据库 Outbox 起步；RocketMQ 可选", "规模和消费者数量上升后接入"),
        ("文件", "S3/MinIO", "附件与证据，数据库保存引用"),
        ("观测", "OpenTelemetry + Prometheus/Grafana", "可先接现有企业监控平台"),
        ("交付", "Docker Compose / 容器平台", "首期不强制 Kubernetes"),
    ], [1700, 3200, 4460], font_size=8.7)
    add_heading(doc, "13.2 工程规范", 2)
    for item in [
        "API contract、domain、infrastructure 分层依赖单向；使用 ArchUnit 防止越层。",
        "数据库变更使用 Flyway/Liquibase；禁止应用启动时无审查自动改生产表。",
        "公共 API、Dubbo contract 和事件 Schema 都遵循向后兼容原则。",
        "领域事件使用过去式命名；命令使用动词；状态枚举禁止复用显示文案。",
        "所有时间以 UTC 存储，界面按租户时区展示；金额使用 decimal + currency。",
        "功能开关用于模板模拟、批量审批和 AI 试点，支持按租户灰度。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. 测试策略", 1)
    add_table(doc, ["测试层级", "覆盖重点", "建议工具 / 做法"], [
        ("单元测试", "聚合不变量、状态转换、规则边界、会签计算", "JUnit 5、属性测试"),
        ("集成测试", "数据库事务、Flowable Adapter、Outbox、租户过滤", "Testcontainers"),
        ("契约测试", "REST、Dubbo DTO、Webhook/事件向后兼容", "OpenAPI 校验、Pact/Schema Test"),
        ("流程测试", "每个模板的成功、拒绝、撤回、超时、并发路径", "流程用例矩阵、时钟可控"),
        ("安全测试", "越权、跨租户、重放、SSRF、附件、敏感日志", "SAST/DAST + 手工审查"),
        ("性能测试", "发起、任务完成、待办查询、Outbox 积压恢复", "k6/JMeter，基于目标 SLI"),
        ("恢复测试", "进程崩溃、重复消息、DB 切换、Webhook 失败", "故障注入与恢复演练"),
    ], [1800, 4400, 3160], font_size=8.9)
    add_heading(doc, "14.1 必测并发场景", 2)
    for item in [
        "同一任务被两个审批人同时提交。",
        "客户端超时后使用相同 Idempotency-Key 重试发起。",
        "本地事务已提交但进程在消息发布前崩溃。",
        "Webhook 接收方处理成功但响应超时，平台再次投递。",
        "模板发布期间有新实例启动，必须只选择一个确定版本。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "15. 实施计划", 1)
    add_heading(doc, "15.1 10-12 周建议里程碑", 2)
    add_table(doc, ["阶段", "周期", "主要交付"], [
        ("M0 方案与骨架", "第 1 周", "领域词汇、仓库结构、CI、数据库基线、OpenAPI/Dubbo contract"),
        ("M1 定义与运行时", "第 2-4 周", "模板/版本、表单校验、Flowable Adapter、实例启动和查询"),
        ("M2 任务与规则", "第 5-7 周", "待办、同意/拒绝/撤回/转交、CEL、审批矩阵、会签"),
        ("M3 集成与安全", "第 8-9 周", "OAuth2、多租户、Webhook、Outbox、审计、附件"),
        ("M4 运营与验收", "第 10-11 周", "管理查询、监控告警、性能/安全/恢复测试、文档"),
        ("缓冲与试点", "第 12 周", "接入 1-2 个真实业务系统、问题修复、上线评审"),
    ], [1900, 1500, 5960])
    add_heading(doc, "15.2 团队配置建议", 2)
    add_body(doc, "基础配置建议为：1 名产品/业务分析、1 名技术负责人、2-3 名后端、1 名前端、1 名测试（可共享）、0.5 名 DevOps/安全支持。若只有 2-3 名研发，应将表单设计器、委托、批量审批和 AI 试点降为 P1。")
    add_heading(doc, "15.3 任务拆分原则", 2)
    for item in [
        "按可验收业务能力拆分，不按 Controller/Service/DAO 技术层拆分迭代。",
        "每个能力同时包含 API、领域逻辑、审计、监控、测试和失败处理。",
        "优先完成一个端到端竖切：发起 → 人工审批 → 完成 → Webhook，再扩展节点类型。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "16. 验收标准", 1)
    add_heading(doc, "16.1 产品验收", 2)
    for item in [
        "至少两个不同业务类型的模板能够独立配置并运行，不需要修改核心代码。",
        "外部系统可通过 API 发起、查询并接收最终结果 Webhook。",
        "支持串行、条件、并行/会签路径以及同意、拒绝、撤回、转交。",
        "管理员能查看完整时间线、规则命中和回调交付记录。",
        "同一发布版本下，使用相同输入得到确定的路径和审批人解析结果。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "16.2 技术验收", 2)
    for item in [
        "通过跨租户访问、重复请求、并发审批和回调重放测试。",
        "达到第 2.3 节约定的性能和可用性指标，或形成经批准的差异记录。",
        "数据库恢复、Outbox 恢复和 Webhook 重放至少完成一次演练。",
        "核心领域模块不依赖 Flowable、Dubbo、ORM 具体实现。",
        "API、事件 Schema、部署、运维、备份和故障处理文档齐全。",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "上线门槛", "若幂等、跨租户隔离、审计完整性、任务并发控制或 Webhook 重放任一项未通过，则不建议进入生产试点。", fill=RED, color="9B1C1C")

    add_heading(doc, "17. 风险与应对", 1)
    add_table(doc, ["风险", "影响", "应对措施"], [
        ("把 DDD 模块过早拆成多个服务", "开发和运维成本过高", "第一阶段只保留两个部署单元；以拆分阈值驱动演进"),
        ("领域模型被 Flowable 对象侵入", "未来升级/替换引擎困难", "使用 WorkflowEnginePort 和映射表，公共模型使用自有 ID"),
        ("审批人解析依赖组织系统", "组织接口故障阻塞发起", "适配器、超时、缓存快照和明确失败策略"),
        ("Webhook 重复导致业务副作用", "重复付款/重复操作", "事件 ID、HMAC、接收方幂等、重放审计"),
        ("规则配置错误", "错误自动通过或错误路径", "发布校验、模拟、版本、灰度、异常默认转人工"),
        ("多租户条件遗漏", "严重数据泄露", "Repository 强制 TenantContext、集成测试、数据库策略评估"),
        ("范围持续膨胀", "MVP 无法按期上线", "P0/P1 门禁，真实业务试点优先，延期能力不进入主路径"),
    ], [2550, 2550, 4260], font_size=8.7)

    add_heading(doc, "18. 后续演进", 1)
    add_heading(doc, "18.1 拆分微服务的触发条件", 2)
    for item in [
        "任务查询与运行时写入出现明显不同的扩容需求。",
        "Webhook/通知故障需要独立隔离和发布。",
        "不同团队分别拥有流程定义、运行时或集成能力。",
        "单体发布频率成为团队协作瓶颈。",
        "某模块需要独立数据生命周期或合规边界。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "18.2 建议演进顺序", 2)
    add_number(doc, "先将 worker 独立部署，并接入 RocketMQ/Kafka。")
    add_number(doc, "按读写压力拆出 Task Query/Search，必要时接入 OpenSearch。")
    add_number(doc, "将 Integration/Notification 拆为独立服务，形成连接器体系。")
    add_number(doc, "建设规则发布治理、脚本沙箱和完整流程测试平台。")
    add_number(doc, "引入 AI 摘要、材料检查、路由建议和低风险自动化评测。")
    add_number(doc, "达到规模与监管要求后，再评估单元化、跨区域容灾和 OceanBase。")

    add_heading(doc, "附录 A：发起审批接口示例", 1)
    add_code_block(doc, """POST /v1/approval-instances
Authorization: Bearer <token>
Idempotency-Key: 01J5-ORDER-20260815-0001
Content-Type: application/json

{
  "templateKey": "purchase_request",
  "externalBusinessKey": "PO-2026-000123",
  "initiator": { "type": "USER", "id": "u_10086" },
  "variables": {
    "amount": 23500.00,
    "currency": "MYR",
    "departmentId": "d_finance",
    "riskLevel": "MEDIUM"
  },
  "callbackContext": { "source": "procurement" }
}""")
    add_heading(doc, "建议响应", 2)
    add_code_block(doc, """HTTP/1.1 201 Created
{
  "requestId": "req_01J5...",
  "instanceId": "appr_01J5...",
  "externalBusinessKey": "PO-2026-000123",
  "status": "RUNNING",
  "definitionVersion": 3,
  "createdAt": "2026-08-15T06:30:00Z"
}""")

    add_heading(doc, "附录 B：状态模型", 1)
    add_table(doc, ["对象", "状态", "说明"], [
        ("ApprovalInstance", "DRAFT", "仅用于允许草稿发起的场景；首期可选"),
        ("ApprovalInstance", "RUNNING", "流程正在等待任务、规则或定时器"),
        ("ApprovalInstance", "APPROVED", "流程按定义成功通过"),
        ("ApprovalInstance", "REJECTED", "被拒绝并按定义结束"),
        ("ApprovalInstance", "WITHDRAWN", "发起人在允许条件下撤回"),
        ("ApprovalInstance", "TERMINATED", "管理员或业务系统强制终止"),
        ("ApprovalTask", "PENDING", "待认领或待处理"),
        ("ApprovalTask", "CLAIMED", "已认领但尚未完成"),
        ("ApprovalTask", "APPROVED/REJECTED", "以审批动作完成"),
        ("ApprovalTask", "TRANSFERRED/CANCELLED", "已转交或因流程推进被取消"),
    ], [2100, 2600, 4660])
    add_body(doc, "状态转换必须通过领域命令完成。普通业务代码和运维脚本不得直接更新状态字段；任何管理员修复都应形成单独命令和审计记录。")

    add_heading(doc, "附录 C：待确认事项与参考资料", 1)
    add_heading(doc, "C.1 立项评审前待确认", 2)
    add_table(doc, ["事项", "建议默认值", "负责人"], [
        ("首个试点场景", "采购申请或费用申请，选择规则清晰、外部副作用可控的场景", "产品/业务"),
        ("数据库", "优先使用团队已有 PostgreSQL 或 MySQL 标准", "架构/运维"),
        ("身份来源", "企业 OIDC/SSO；组织关系通过适配器读取", "安全/IT"),
        ("消息队列", "MVP 可用 Outbox 轮询；已有 RocketMQ 时直接接入", "架构"),
        ("表单设计器", "首期基础组件，不建设通用低代码平台", "产品/前端"),
        ("RPO/RTO 与保留期", "RPO 15 分钟、RTO 2 小时；保留期按业务合规确认", "业务/合规/运维"),
    ], [2600, 5200, 1560], font_size=8.8)
    add_heading(doc, "C.2 参考资料", 2)
    refs = [
        "Apache Dubbo Protocol Overview: https://dubbo.apache.org/en/overview/mannual/java-sdk/reference-manual/protocol/overview/",
        "Flowable Open Source Documentation: https://www.flowable.com/open-source/docs/oss-introduction",
        "Apache Seata Transaction Models: https://seata.apache.org/docs/v2.5/overview/what-is-seata/",
        "Common Expression Language: https://cel.dev/",
        "OpenTelemetry Java: https://opentelemetry.io/docs/languages/java/",
        "NIST AI Risk Management Framework: https://airc.nist.gov/airmf-resources/airmf/",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    add_heading(doc, "C.3 文档结论", 2)
    add_callout(doc, "推荐立项基线", "以模块化 DDD 核心服务为主体，用 Dubbo Triple 提供内部契约，用 REST/Webhook 服务外部系统，用 Flowable 承担可靠流程状态机。第一阶段先做可上线的通用审批闭环，再依据真实压力和组织边界拆分微服务。", fill=LIGHT_BLUE)

    doc.core_properties.title = "通用审批流微服务第一阶段技术方案"
    doc.core_properties.subject = "轻量级通用审批平台 MVP 架构与实施方案"
    doc.core_properties.author = "Approval Platform Architecture Team"
    doc.core_properties.keywords = "审批流, Dubbo, DDD, Flowable, 微服务, MVP"
    doc.core_properties.comments = "Generated as a first-phase solution baseline."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("DOCX created")


if __name__ == "__main__":
    build_document()
